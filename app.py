import streamlit as st
import torch
from transformers import pipeline
import docx
from pypdf import PdfReader
import io
from docx import Document

# --- PAGE CONFIG ---
st.set_page_config(page_title="Muzi AI Quiz Studio Pro", page_icon="🎯", layout="wide")

# Professional UI Styling
st.markdown("""
    <style>
    .stButton>button { width: 100%; border-radius: 8px; background-color: #1e3c72; color: white; height: 3em; font-weight: bold; }
    .quiz-container { 
        padding: 20px; border-radius: 10px; border-left: 8px solid #1e3c72; 
        background-color: #ffffff !important; color: #1a1a1a !important; 
        margin-bottom: 20px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    .quiz-container b { color: #1e3c72 !important; }
    .quiz-container p { color: #1a1a1a !important; white-space: pre-wrap; font-family: sans-serif; }
    </style>
    """, unsafe_allow_html=True)

# --- MODEL LOADING (STABLE) ---
@st.cache_resource
def load_stable_model():
    return pipeline("text-generation", model="MBZUAI/LaMini-GPT-124M")

generator = load_stable_model()

# --- HELPERS ---
def extract_text(file):
    if file.name.endswith('.pdf'):
        reader = PdfReader(file)
        return "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
    elif file.name.endswith('.docx'):
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

def create_docx(quizzes):
    doc = Document()
    doc.add_heading('🎯 Muzi AI Quiz Studio Pro Report', 0)
    for i, q in enumerate(quizzes, 1):
        doc.add_heading(f'Quiz Version {i}', level=1)
        doc.add_paragraph(q)
        doc.add_page_break()
    return doc

# --- UI ---
st.title("🎯 Muzi AI Quiz Studio PRO")
st.write("NUST Balochistan Campus - Professional Assessment Tool")

# Session state to store quizzes
if 'quizzes' not in st.session_state:
    st.session_state.quizzes = []

with st.sidebar:
    st.header("⚙️ Configuration")
    uploaded_file = st.file_uploader("Upload Document", type=['pdf', 'docx'])
    st.divider()
    # ALL OLD OPTIONS RESTORED
    num_versions = st.slider("Total Quizzes", 1, 20, 1)
    q_per_quiz = st.slider("MCQs per Quiz", 1, 20, 5)
    difficulty = st.selectbox("Difficulty:", ["Easy", "Standard", "Advanced"])
    st.info("Stable mode enabled to prevent crashes.")

if st.button("🚀 GENERATE PROFESSIONAL MCQS"):
    if uploaded_file:
        context_raw = extract_text(uploaded_file)
        context = context_raw[:600] # Keeping it short for model focus
        st.session_state.quizzes = []
        
        for i in range(1, num_versions + 1):
            full_quiz = ""
            with st.status(f"Generating Quiz {i}...") as status:
                for j in range(1, q_per_quiz + 1):
                    prompt = f"Context: {context}\n\nTask: Create one {difficulty} MCQ with options A, B, C, D and the answer.\nOutput:"
                    
                    output = generator(prompt, max_length=256, do_sample=True, temperature=0.7)
                    res = output[0]['generated_text'].split("Output:")[-1].strip()
                    
                    full_quiz += f"Question {j}: {res}\n\n"
                
                st.session_state.quizzes.append(full_quiz)
                st.markdown(f"<div class='quiz-container'><b>📝 VERSION {i}</b><p>{full_quiz}</p></div>", unsafe_allow_html=True)
                status.update(label=f"Quiz {i} Complete!", state="complete")
    else:
        st.error("Muzi bhai, pehle file upload karein!")

# DOWNLOAD BUTTON RESTORED
if st.session_state.quizzes:
    doc = create_docx(st.session_state.quizzes)
    bio = io.BytesIO()
    doc.save(bio)
    st.download_button("📥 DOWNLOAD ALL (WORD FILE)", data=bio.getvalue(), file_name="Muzi_Pro_Quizzes.docx")
